#!/usr/bin/env python3
"""
Student Worksheet Converter - Markdown to DOCX

This utility converts student worksheets from Markdown format to Microsoft Word (DOCX) format
for easy distribution and printing. It maintains formatting, adds proper styling, and handles
educational content appropriately.

Usage:
    python src/worksheet_converter.py [options]

Examples:
    # Convert all worksheets
    python src/worksheet_converter.py --all

    # Convert specific worksheet
    python src/worksheet_converter.py --file sast-student-worksheet.md

    # Convert with custom output directory
    python src/worksheet_converter.py --all --output-dir ./converted_worksheets
"""

import os
import sys
import argparse
import re
from pathlib import Path
from typing import List, Dict, Any

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    import markdown
    from markdown.extensions import codehilite, tables
except ImportError as e:
    print(f"❌ Missing required dependencies: {e}")
    print("📦 Install with: pip install python-docx markdown")
    sys.exit(1)


class WorksheetConverter:
    """Converts Markdown student worksheets to DOCX format."""

    def __init__(self, verbose: bool = False):
        self.verbose = verbose
        self.worksheets_dir = Path("docs/student-worksheets")
        self.converted_count = 0

    def log(self, message: str) -> None:
        """Print verbose log messages."""
        if self.verbose:
            print(f"🔧 {message}")

    def setup_document_styles(self, doc: Document) -> None:
        """Configure document styles for educational worksheets."""
        # Document settings
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Heading styles
        try:
            heading1 = doc.styles['Heading 1']
            heading1.font.size = Pt(16)
            heading1.font.bold = True
            heading1.font.color.rgb = None  # Black

            heading2 = doc.styles['Heading 2']
            heading2.font.size = Pt(14)
            heading2.font.bold = True

            heading3 = doc.styles['Heading 3']
            heading3.font.size = Pt(12)
            heading3.font.bold = True

        except KeyError:
            self.log("Warning: Some default styles not available")

    def add_student_info_header(self, doc: Document) -> None:
        """Add student information header to the document."""
        # Title
        title = doc.add_heading('', level=1)
        title_run = title.runs[0] if title.runs else title.add_run()
        title_run.font.size = Pt(18)
        title_run.font.bold = True

        # Student info table
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'

        # First row
        cells = table.rows[0].cells
        cells[0].text = "Name: _________________________________"
        cells[1].text = "Date: ________________"

        # Second row
        cells = table.rows[1].cells
        cells[0].text = "Lab Partner: __________________________"
        cells[1].text = "Section: ______________"

        # Add spacing
        doc.add_paragraph("")

    def convert_markdown_content(self, content: str) -> str:
        """Convert Markdown to HTML, handling educational content specially."""
        # Pre-process content for better DOCX conversion

        # Convert checkboxes to proper format
        content = re.sub(r'- \[ \]', '☐', content)
        content = re.sub(r'- \[x\]', '☑', content)

        # Handle fill-in blanks (underscores)
        # Standardize blank lengths
        content = re.sub(r'_{10,}', '_' * 30, content)

        # Convert to HTML
        md = markdown.Markdown(
            extensions=['tables', 'codehilite', 'fenced_code'])
        html_content = md.convert(content)

        return html_content

    def add_markdown_to_docx(self, doc: Document, markdown_content: str) -> None:
        """Add Markdown content to DOCX document with appropriate formatting."""
        lines = markdown_content.split('\n')
        current_section = None
        in_code_block = False
        code_language = None

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Skip empty lines in most cases
            if not line:
                if not in_code_block:
                    doc.add_paragraph("")
                i += 1
                continue

            # Handle code blocks
            if line.startswith('```'):
                if not in_code_block:
                    # Starting code block
                    in_code_block = True
                    code_language = line[3:].strip() if len(
                        line) > 3 else 'text'
                    self.log(f"Starting code block: {code_language}")
                else:
                    # Ending code block
                    in_code_block = False
                    code_language = None
                    self.log("Ending code block")
                i += 1
                continue

            # Handle content inside code blocks
            if in_code_block:
                p = doc.add_paragraph(line)
                p.style = 'No Spacing'
                # Make code text monospace and slightly smaller
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                i += 1
                continue

            # Handle headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()

                # Remove emojis for cleaner Word document
                header_text = re.sub(
                    r'[🎯📋🔧🛠️📚🧠🔍⚖️💡🚀📊🎓]', '', header_text).strip()

                if level <= 3:
                    doc.add_heading(header_text, level=level)
                else:
                    p = doc.add_paragraph(header_text)
                    p.style = 'Heading 3'

                i += 1
                continue

            # Handle tables
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_lines = []
                j = i
                while j < len(lines) and '|' in lines[j]:
                    table_lines.append(lines[j])
                    j += 1

                if len(table_lines) >= 2:
                    self.add_table_to_docx(doc, table_lines)
                    i = j
                    continue

            # Handle lists
            if line.startswith(('-', '*', '+')) or re.match(r'^\d+\.', line):
                self.add_list_item_to_docx(doc, line)
                i += 1
                continue

            # Handle regular paragraphs
            if line:
                # Process inline formatting
                formatted_text = self.process_inline_formatting(line)
                p = doc.add_paragraph(formatted_text)

                # Special formatting for questions and fill-ins
                if ':' in line and ('?' in line or 'Fill in' in line):
                    p.style = 'Intense Quote'

            i += 1

    def add_table_to_docx(self, doc: Document, table_lines: List[str]) -> None:
        """Add a table from Markdown format to the DOCX document."""
        # Parse table
        rows = []
        for line in table_lines:
            if '---' in line or '===' in line:
                continue  # Skip separator lines
            cells = [cell.strip() for cell in line.split(
                '|')[1:-1]]  # Remove empty first/last
            if cells:
                rows.append(cells)

        if not rows:
            return

        # Create table
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'

        # Fill table
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = cell_data

                    # Make header row bold
                    if i == 0:
                        for paragraph in row.cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

    def add_list_item_to_docx(self, doc: Document, line: str) -> None:
        """Add a list item to the DOCX document."""
        # Remove list markers
        text = re.sub(r'^[-*+]\s*', '', line)
        text = re.sub(r'^\d+\.\s*', '', text)

        p = doc.add_paragraph(text, style='List Bullet')

        # Handle checkboxes
        if '☐' in text or '☑' in text:
            p.style = 'List Bullet'

    def process_inline_formatting(self, text: str) -> str:
        """Process inline Markdown formatting."""
        # Remove emojis for cleaner Word document
        text = re.sub(r'[🎯📋🔧🛠️📚🧠🔍⚖️💡🚀📊🎓✅❌⚠️]', '', text)

        # Convert bold/italic (basic - DOCX will handle this differently)
        # Remove markdown, let docx handle formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)

        # Handle code spans
        text = re.sub(r'`(.*?)`', r'\1', text)  # Remove backticks

        return text.strip()

    def convert_worksheet(self, markdown_file: Path, output_dir: Path) -> Path:
        """Convert a single worksheet from Markdown to DOCX."""
        self.log(f"Converting {markdown_file.name}")

        # Read Markdown content
        try:
            with open(markdown_file, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception as e:
            print(f"❌ Error reading {markdown_file}: {e}")
            return None

        # Create new document
        doc = Document()
        self.setup_document_styles(doc)

        # Extract title from first line
        first_line = content.split('\n')[0].strip()
        if first_line.startswith('#'):
            title = first_line.lstrip('#').strip()
            title = re.sub(r'[🎯📋🔧🛠️📚🧠🔍⚖️💡🚀📊🎓]', '', title).strip()
        else:
            title = markdown_file.stem.replace('-', ' ').title()

        # Add title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add student info header
        self.add_student_info_header(doc)

        # Add horizontal line
        doc.add_paragraph("─" * 80)

        # Process the rest of the content (skip the title line)
        content_lines = content.split('\n')[1:]  # Skip title
        remaining_content = '\n'.join(content_lines)

        # Add main content
        self.add_markdown_to_docx(doc, remaining_content)

        # Determine output file path
        output_file = output_dir / f"{markdown_file.stem}.docx"

        # Save document
        try:
            doc.save(str(output_file))
            self.converted_count += 1
            print(f"✅ Converted: {markdown_file.name} → {output_file.name}")
            return output_file
        except Exception as e:
            print(f"❌ Error saving {output_file}: {e}")
            return None

    def convert_all_worksheets(self, output_dir: Path = None) -> List[Path]:
        """Convert all student worksheets to DOCX format."""
        if output_dir is None:
            output_dir = self.worksheets_dir

        # Ensure output directory exists
        output_dir.mkdir(parents=True, exist_ok=True)

        # Find all worksheet files
        worksheet_files = list(self.worksheets_dir.glob("*.md"))

        if not worksheet_files:
            print(f"❌ No worksheet files found in {self.worksheets_dir}")
            return []

        print(f"📄 Found {len(worksheet_files)} worksheet files")

        converted_files = []
        for worksheet_file in worksheet_files:
            output_file = self.convert_worksheet(worksheet_file, output_dir)
            if output_file:
                converted_files.append(output_file)

        return converted_files

    def convert_single_worksheet(self, filename: str, output_dir: Path = None) -> Path:
        """Convert a single worksheet by filename."""
        if output_dir is None:
            output_dir = self.worksheets_dir

        # Ensure output directory exists
        output_dir.mkdir(parents=True, exist_ok=True)

        worksheet_file = self.worksheets_dir / filename

        if not worksheet_file.exists():
            print(f"❌ Worksheet file not found: {worksheet_file}")
            return None

        return self.convert_worksheet(worksheet_file, output_dir)


def main():
    """Main function for command-line usage."""
    parser = argparse.ArgumentParser(
        description="Convert student worksheets from Markdown to DOCX format",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s --all                              # Convert all worksheets
  %(prog)s --file sast-student-worksheet.md  # Convert specific worksheet
  %(prog)s --all --output-dir ./converted     # Convert to custom directory
        """
    )

    parser.add_argument(
        '--all',
        action='store_true',
        help='Convert all student worksheets'
    )

    parser.add_argument(
        '--file',
        type=str,
        help='Convert specific worksheet file (e.g., sast-student-worksheet.md)'
    )

    parser.add_argument(
        '--output-dir',
        type=str,
        help='Output directory for DOCX files (default: same as source)'
    )

    parser.add_argument(
        '--verbose', '-v',
        action='store_true',
        help='Enable verbose output'
    )

    parser.add_argument(
        '--list',
        action='store_true',
        help='List available worksheet files'
    )

    args = parser.parse_args()

    # Change to project root directory
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    os.chdir(project_root)

    converter = WorksheetConverter(verbose=args.verbose)

    # List worksheets
    if args.list:
        worksheet_files = list(converter.worksheets_dir.glob("*.md"))
        print(f"📄 Available worksheets in {converter.worksheets_dir}:")
        for i, file in enumerate(worksheet_files, 1):
            print(f"  {i}. {file.name}")
        return

    # Determine output directory
    output_dir = Path(
        args.output_dir) if args.output_dir else converter.worksheets_dir

    print("📚 Student Worksheet Converter")
    print("=" * 50)

    # Convert worksheets
    if args.all:
        print("🔄 Converting all student worksheets...")
        converted_files = converter.convert_all_worksheets(output_dir)

        print("\n" + "=" * 50)
        print(
            f"✅ Conversion complete! {converter.converted_count} files converted.")
        print(f"📁 Output directory: {output_dir.absolute()}")

        if converted_files:
            print("\n📄 Converted files:")
            for file in converted_files:
                print(f"  • {file.name}")

    elif args.file:
        print(f"🔄 Converting {args.file}...")
        output_file = converter.convert_single_worksheet(args.file, output_dir)

        if output_file:
            print(f"✅ Conversion complete!")
            print(f"📄 Output file: {output_file}")
        else:
            print("❌ Conversion failed!")
            sys.exit(1)

    else:
        print("❌ Please specify --all or --file option")
        parser.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()
